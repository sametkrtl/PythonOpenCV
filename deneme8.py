# -*- coding: utf-8 -*-
# WORD TABANLI VERI MASKELEME (NER + SAHTE VERI - GELISMIS CHUNK DESTEKLI SURUM)

import os, uuid, re, tempfile
import gradio as gr
from docx import Document
from transformers import pipeline
from pdf2docx import Converter
from docx2pdf import convert
from textwrap import wrap

# --- SABIT SAHTE VERILER ---
SABIT_VERILER = {
    "ad_soyad": "Ahmet Yılmaz",
    "telefon": "05001234567",
    "adres": "İstanbul, Türkiye",
    "tc_kimlik": "11111111110",
    "tarih": "01.01.2000",
    "para": "1000 TL",
    "sirket": "Örnek A.Ş."
}

# --- SAHTE DEGER SECIMI ---
def sahte_deger_uret(label, uzunluk=None):
    deger = SABIT_VERILER.get(label, "***")
    return deger[:uzunluk] if uzunluk else deger

# --- TC DOGRULAMA ---
def is_valid_tc(tc):
    if not tc.isdigit() or len(tc) != 11 or tc[0] == '0': return False
    digits = list(map(int, tc))
    d10 = ((sum(digits[0:9:2]) * 7) - sum(digits[1:8:2])) % 10
    d11 = sum(digits[:10]) % 10
    return digits[9] == d10 and digits[10] == d11

# --- NER MODEL ---
print("NER modeli yukleniyor...")
ner_pipeline = pipeline("ner", model="./ner_model_final", tokenizer="./ner_model_final", aggregation_strategy="simple")

# --- GELISMIS CHUNK DESTEKLI NER + REGEX ---
def get_combined_entities(text, use_tc):
    chunks = wrap(text, 400, break_long_words=False, break_on_hyphens=False)

    all_model_entities = []
    offset = 0
    for chunk in chunks:
        entities = ner_pipeline(chunk)
        for ent in entities:
            ent["start"] += offset
            ent["end"] += offset
        all_model_entities.extend(entities)
        offset += len(chunk)

    print(f"\n📊 MODEL BULDUĞU {len(all_model_entities)} ETİKET:")
    for ent in all_model_entities:
        print(f"📌 '{ent['word']}' [{ent['entity_group']}] ({ent['start']}–{ent['end']}) Skor: {ent.get('score',1.0):.2f}")

    regex_entities = []
    if use_tc:
        for match in re.finditer(r"\b[1-9][0-9]{10}\b", text):
            if is_valid_tc(match.group(0)):
                regex_entities.append({"entity_group": "tc_kimlik", "start": match.start(), "end": match.end(), "word": match.group(0)})

    all_spans = {(e['start'], e['end']) for e in regex_entities}
    final = regex_entities + [e for e in all_model_entities if not any(max(e['start'], s)<min(e['end'], e_) for (s,e_) in all_spans)]
    return sorted(final, key=lambda x: x['start'])

# --- PDF -> DOCX ---
def pdf_to_docx(pdf_path):
    docx_path = pdf_path.replace(".pdf", ".docx")
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()
    return docx_path

# --- DOCX MASKELEME (KELIME BAZLI YAZMA) ---
def maskele_docx(docx_path, cikti_path, tc_kontrol=True, sahte=True):
    doc = Document(docx_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])
    entities = get_combined_entities(full_text, use_tc=tc_kontrol)
    if not entities:
        doc.save(cikti_path)
        return

    for para in doc.paragraphs:
        for run in para.runs:
            for ent in entities:
                if ent['word'] in run.text:
                    replacement = sahte_deger_uret(ent['entity_group'], len(ent['word'])) if sahte else "*" * len(ent['word'])
                    run.text = run.text.replace(ent['word'], replacement)
    doc.save(cikti_path)

# --- DOCX -> PDF ---
def docx_to_pdf(docx_path):
    pdf_path = docx_path.replace(".docx", "_masked.pdf")
    convert(docx_path, pdf_path)
    return pdf_path

# --- GRADIO ARAYUZ ---
with gr.Blocks(title="Word + PDF Maskeleme") as app:
    gr.Markdown("# 📄 Word Tabanlı Veri Maskeleme (Chunk Destekli)")
    with gr.Tab("PDF -> Word -> PDF Maskeleme"):
        pdf_input = gr.File(label="📎 PDF Yükle", file_types=[".pdf"])
        tc_check = gr.Checkbox(label="🆔 TC Kimlik Doğrulama", value=True)
        sahte_check = gr.Checkbox(label="🎭 Sabit Veri Kullan", value=True)
        out_pdf = gr.File(label="📥 Maske Sonucu (PDF)")
        btn = gr.Button("🚀 Maskele")

    def pipeline_donustur(pdf_dosya, tc_var, sahte_var):
        docx = pdf_to_docx(pdf_dosya.name)
        cikti_docx = os.path.join(tempfile.gettempdir(), f"masked_{uuid.uuid4().hex}.docx")
        maskele_docx(docx, cikti_docx, tc_var, sahte_var)
        return docx_to_pdf(cikti_docx)

    btn.click(fn=pipeline_donustur, inputs=[pdf_input, tc_check, sahte_check], outputs=out_pdf)

if __name__ == "__main__":
    app.launch()
