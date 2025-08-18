import gradio as gr
import os
import time
import json
import pandas as pd
from datetime import datetime
from pathlib import Path
import tempfile
import shutil
from typing import List, Dict, Tuple, Optional
import logging
from faker import Faker
import random
# ----------------------------------------------------
#  EnhancedAnonymizationApp  (tam sürüm)
# ----------------------------------------------------
class EnhancedAnonymizationApp:
    def __init__(self):
        # Logger ayarları
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)
        
        # self.pdf_replacer = PDFTextReplacer()  # Gerçek implementasyonda açın
        # Tutarlı sahte üretim için sabit seed
        self.global_seed = 123456
        random.seed(self.global_seed)

        # 🇹🇷 Türkçe Faker örneği
        self.fake = Faker("tr_TR")
        self.fake.seed_instance(self.global_seed)

        # Dizinleri oluştur
        for dir_name in ["uploads", "outputs", "logs", "temp"]:
            os.makedirs(dir_name, exist_ok=True)
        
        self.processing_stats = []
    
    def process_pdf_with_real_replacement(self, 
                                        pdf_file, 
                                        confidence_threshold: float,
                                        replacement_strategy: str,
                                        custom_replacements: str,
                                        enable_logging: bool,
                                        enable_debug: bool,
                                        conversion_method: str,
                                        progress=gr.Progress()) -> Tuple[Optional[str], str, str, pd.DataFrame, str]:
        """
        Gelişmiş PDF işleme - GERÇEK değiştirme ile
        Dönüş: (output_pdf_path | None, status_msg, summary_md, stats_df, debug_text)
        """
        if pdf_file is None:
            return None, "❌ Lütfen bir PDF dosyası yükleyin.", "", pd.DataFrame(), ""
        
        debug_info = ""
        
        try:
            progress(0.05, desc="Dosya hazırlanıyor...")
            
            # Dosya yolları
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_in = os.path.basename(pdf_file.name if hasattr(pdf_file, "name") else str(pdf_file))
            input_filename = f"input_{timestamp}_{base_in}"
            output_filename = f"anonymized_{timestamp}_{base_in}"
            
            input_path = os.path.join("uploads", input_filename)
            output_path = os.path.join("outputs", output_filename)
            
            # Yüklenen dosyayı kopyala
            shutil.copy2(pdf_file.name if hasattr(pdf_file, "name") else str(pdf_file), input_path)
            
            progress(0.1, desc="PDF analiz ediliyor...")
            
            # 1. PDF'ten metin çıkar ve analiz et
            entities_detected = self.extract_and_analyze_pdf(input_path, confidence_threshold, progress)
            
            if not entities_detected:
                return None, "⚠️ PDF'de kişisel bilgi tespit edilmedi.", "", pd.DataFrame(), "Hiçbir entity tespit edilmedi."
            
            progress(0.4, desc="Değiştirme stratejisi uygulanıyor...")
            
            # 2. Değiştirme stratejisini uygula
            processed_entities = self.apply_replacement_strategy(
                entities_detected, 
                replacement_strategy, 
                custom_replacements or ""
            )
            
            progress(0.5, desc="PDF'te değişiklikler uygulanıyor...")
            
            # 3. Gerçek PDF değiştirme
            success = self.perform_actual_pdf_replacement(
                input_path, 
                processed_entities, 
                output_path, 
                conversion_method,
                progress
            )
            
            if not success:
                return None, "❌ PDF değiştirme işlemi başarısız oldu.", "", pd.DataFrame(), "Değiştirme başarısız"
            
            progress(0.8, desc="Sonuçlar doğrulanıyor...")
            
            # 4. Değiştirmeleri doğrula
            validation_result = self.validate_pdf_changes(output_path, processed_entities)
            debug_info = self.format_debug_info(validation_result, processed_entities)
            
            progress(0.9, desc="İstatistikler oluşturuluyor...")
            
            # 5. İstatistikleri oluştur
            stats_df = self.create_detailed_statistics(processed_entities, validation_result)
            summary = self.generate_detailed_summary(processed_entities, validation_result, input_filename)
            
            # 6. Loglama
            if enable_logging:
                self.log_processing_with_validation(input_filename, processed_entities, validation_result)
            
            progress(1.0, desc="Tamamlandı!")
            
            # Başarı mesajı
            success_rate = validation_result.get('successfully_replaced', 0)
            total_entities = len(processed_entities)
            status_msg = f"✅ İşlem tamamlandı! {success_rate}/{total_entities} değiştirme başarılı."
            
            return output_path, status_msg, summary, stats_df, (debug_info if enable_debug else "")
            
        except Exception as e:
            self.logger.error(f"PDF işleme hatası: {e}", exc_info=True)
            error_msg = f"❌ Kritik hata: {str(e)}"
            return None, error_msg, "", pd.DataFrame(), f"Hata detayı: {str(e)}"
    
    def extract_and_analyze_pdf(self, pdf_path: str, confidence_threshold: float, progress) -> List[Dict]:
        """PDF'i analiz et ve entity'leri tespit et"""
        try:
            import fitz  # PyMuPDF
            
            # PDF'i aç ve metni çıkar
            doc = fitz.open(pdf_path)
            full_text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                full_text += page_text + "\n"
            
            doc.close()
            
            progress(0.2, desc=f"Metin çıkarıldı ({len(full_text)} karakter)...")
            
            if len(full_text.strip()) < 10:
                self.logger.warning("PDF'den yeterli metin çıkarılamadı")
                return []
            
            progress(0.3, desc="NER analizi yapılıyor...")
            
            # NER ile entity tespiti (mock; yerine kendi modelinizi entegre edin)
            entities = self.perform_ner_analysis(full_text, confidence_threshold)
            
            # Regex ile ek tespitler
            regex_entities = self.perform_regex_detection(full_text)
            
            # Birleştir ve temizle
            combined_entities = self.merge_and_clean_entities(entities + regex_entities, confidence_threshold)
            
            self.logger.info(f"Toplam {len(combined_entities)} entity tespit edildi")
            return combined_entities
            
        except Exception as e:
            self.logger.error(f"PDF analiz hatası: {e}")
            return []
    
    def perform_ner_analysis(self, text: str, confidence_threshold: float) -> List[Dict]:
        """NER analizi (mock)"""
        import re
        mock_entities = []
        
        patterns = {
            'ad_soyad': r'\b[A-ZÇĞÖŞÜİ][a-zçğöşüı]+\s+[A-ZÇĞÖŞÜİ][a-zçğöşüı]+\b',
            'tc_kimlik': r'\b[1-9][0-9]{9}[02468]\b',
            'telefon': r'\b0?5[0-9]{2}[\s\-]?[0-9]{3}[\s\-]?[0-9]{2}[\s\-]?[0-9]{2}\b',
            'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'para': r'\b\d{1,3}(?:\.\d{3})*(?:,\d{2})?\s?(?:TL|₺|lira)\b'
        }
        
        for entity_type, pattern in patterns.items():
            matches = re.finditer(pattern, text)
            for match in matches:
                mock_entities.append({
                    'entity': entity_type,
                    'word': match.group(),
                    'start': match.start(),
                    'end': match.end(),
                    'score': 0.8 + (hash(match.group()) % 20) / 100  # Mock confidence
                })
        
        return [e for e in mock_entities if e['score'] >= confidence_threshold]
    
    def perform_regex_detection(self, text: str) -> List[Dict]:
        """Regex ile ek tespitler (TC kimlik validasyonu)"""
        import re
        regex_entities = []
        
        tc_pattern = r'\b[1-9][0-9]{9}[02468]\b'
        for match in re.finditer(tc_pattern, text):
            tc_no = match.group()
            if self.validate_turkish_id(tc_no):
                regex_entities.append({
                    'entity': 'tc_kimlik',
                    'word': tc_no,
                    'start': match.start(),
                    'end': match.end(),
                    'score': 0.95,
                    'method': 'regex_validated'
                })
        
        return regex_entities
    
    def validate_turkish_id(self, tc_no: str) -> bool:
        """TC kimlik validasyonu"""
        if not tc_no or len(tc_no) != 11 or not tc_no.isdigit():
            return False
        if tc_no[0] == '0':
            return False
        digits = [int(d) for d in tc_no]
        odd_sum = sum(digits[i] for i in range(0, 9, 2))
        even_sum = sum(digits[i] for i in range(1, 8, 2))
        check_digit_10 = (odd_sum * 7 - even_sum) % 10
        if check_digit_10 != digits[9]:
            return False
        check_digit_11 = sum(digits[:10]) % 10
        if check_digit_11 != digits[10]:
            return False
        return True
    
    def merge_and_clean_entities(self, entities: List[Dict], confidence_threshold: float) -> List[Dict]:
        """Entity'leri birleştir ve temizle"""
        seen = set()
        cleaned = []
        for entity in entities:
            key = (entity.get('start', 0), entity.get('end', 0), entity.get('word', ''))
            if key not in seen:
                seen.add(key)
                cleaned.append(entity)
        filtered = [e for e in cleaned if e.get('score', 0) >= confidence_threshold]
        return sorted(filtered, key=lambda x: x.get('start', 0))
    
    # ------------ Faker yardımcıları ------------
    def _seeded_faker(self, key: str) -> Faker:
        """Aynı 'key' için deterministik Faker."""
        seeded = abs(hash(key)) % (10**9)
        f = Faker("tr_TR")
        f.seed_instance(seeded)
        return f

    def _fake_tc_kimlik(self, f: Faker) -> str:
        """Geçerli checksum'lı TR TCKN üret."""
        while True:
            d = [0]*11
            d[0] = f.random_int(1, 9)
            for i in range(1, 9):
                d[i] = f.random_int(0, 9)
            odd_sum = sum(d[i] for i in range(0, 9, 2))
            even_sum = sum(d[i] for i in range(1, 8, 2))
            d[9] = (odd_sum * 7 - even_sum) % 10
            d[10] = (sum(d[:10]) % 10)
            tc = "".join(map(str, d))
            if self.validate_turkish_id(tc):
                return tc

    def _fake_phone_tr(self, f: Faker) -> str:
        """05xxxxxxxxx formatında cep."""
        return "05" + "".join(str(f.random_int(0, 9)) for _ in range(9))

    def _fake_iban_tr(self, f: Faker) -> str:
        try:
            return f.iban(country_code="TR")
        except Exception:
            return "TR" + "".join(str(f.random_int(0, 9)) for _ in range(24))

    def _fake_company_tr(self, f: Faker) -> str:
        try:
            return f.company()
        except Exception:
            return "Örnek A.Ş."

    def _fake_address_tr(self, f: Faker) -> str:
        try:
            return f.address().replace("\n", ", ")
        except Exception:
            return "Örnek Mah., Örnek Cd., No:1, İstanbul"

    def _fake_name_tr(self, f: Faker) -> str:
        try:
            return f.name()
        except Exception:
            return "Ali Demir"

    def _fake_email_from_name(self, f: Faker, name_full: str) -> str:
        import unicodedata, re
        def slugify(s: str) -> str:
            s = unicodedata.normalize("NFKD", s)
            s = "".join(c for c in s if not unicodedata.combining(c))
            s = s.lower()
            s = (s.replace("ı","i").replace("ğ","g").replace("ş","s")
                   .replace("ç","c").replace("ö","o").replace("ü","u"))
            s = re.sub(r"[^a-z0-9]+", ".", s).strip(".")
            return s

        parts = (name_full or "").split()
        if len(parts) >= 2:
            user = f"{slugify(parts[0])}.{slugify(parts[-1])}"
        else:
            user = slugify(name_full or "kullanici")
        domain = f.random_element(["example.com","mail.com","ornek.com","kurum.com.tr"])
        return f"{user}@{domain}"
    
    def apply_replacement_strategy(self, entities: List[Dict], strategy: str, custom_replacements: str) -> List[Dict]:
        """Değiştirme stratejisini uygula (Faker destekli)."""
        default_replacements = {
            'ad_soyad': 'Ali Demir',
            'tc_kimlik': '11111111110',
            'telefon': '05009999999',
            'adres': 'Örnek Mahallesi, İstanbul',
            'para': '1000 TL',
            'tarih': '01.01.2000',
            'email': 'ornek@email.com',
            'sirket': 'Örnek A.Ş.',
            'iban': 'TR00 0000 0000 0000 0000 0000 00'
        }
        if custom_replacements and custom_replacements.strip():
            try:
                custom_dict = json.loads(custom_replacements)
                default_replacements.update(custom_dict)
            except json.JSONDecodeError as e:
                self.logger.warning(f"Custom replacements parse hatası: {e}")

        processed = []
        for entity in entities:
            e = entity.copy()
            etype = e.get('entity', '').strip()
            original = e.get('word', '')

            if strategy == "Maskeleme (*)":
                e['replacement'] = '*' * len(original)

            elif strategy == "Genel Değerler":
                generic_map = {
                    'ad_soyad': '[İSİM]',
                    'tc_kimlik': '[TC KİMLİK]',
                    'telefon': '[TELEFON]',
                    'adres': '[ADRES]',
                    'para': '[PARA]',
                    'tarih': '[TARİH]',
                    'email': '[E-POSTA]',
                    'sirket': '[ŞİRKET]',
                    'iban': '[IBAN]'
                }
                e['replacement'] = generic_map.get(etype, f'[{etype.upper()}]')

            elif strategy == "Sahte Değerler":
                # Aynı orijinal -> aynı sahte değer (deterministik)
                f = self._seeded_faker(original if original else etype)

                # Kullanıcının JSON ile verdiği tür varsa öncelik
                override = default_replacements.get(etype)
                candidate = None if override in (None, "") else override

                if candidate is None:
                    if etype == 'ad_soyad':
                        candidate = self._fake_name_tr(f)
                    elif etype == 'telefon':
                        candidate = self._fake_phone_tr(f)
                    elif etype == 'email':
                        base_name = original if original else self._fake_name_tr(f)
                        candidate = self._fake_email_from_name(f, base_name)
                    elif etype == 'adres':
                        candidate = self._fake_address_tr(f)
                    elif etype == 'sirket':
                        candidate = self._fake_company_tr(f)
                    elif etype == 'iban':
                        candidate = self._fake_iban_tr(f)
                    elif etype == 'tc_kimlik':
                        candidate = self._fake_tc_kimlik(f)
                    elif etype == 'para':
                        tl = f.pydecimal(left_digits=4, right_digits=2, positive=True)
                        candidate = f"{tl} TL"
                    elif etype == 'tarih':
                        dt = f.date_between(start_date="-10y", end_date="today")
                        candidate = dt.strftime("%d.%m.%Y")
                    else:
                        # Tanımsız tür: uzunluğu kabaca koruyan fallback
                        candidate = original if original else "—"

                e['replacement'] = candidate

            else:
                e['replacement'] = original

            processed.append(e)
        return processed

    
    def perform_actual_pdf_replacement(self, input_path: str, entities: List[Dict], 
                                     output_path: str, method: str, progress) -> bool:
        """Gerçek PDF değiştirme işlemi"""
        try:
            progress(0.6, desc=f"{method} yöntemiyle değiştiriliyor...")
            
            if method in ("DOCX Yöntemi", "auto"):
                success = self._replace_via_docx_method(input_path, entities, output_path, progress)
                if success:
                    return True
            
            if method in ("PyMuPDF Direkt", "auto"):
                progress(0.65, desc="PyMuPDF ile direkt değiştiriliyor...")
                success = self._replace_via_pymupdf_method(input_path, entities, output_path)
                if success:
                    return True
            
            if method in ("Text Recreation", "auto"):
                progress(0.7, desc="Metin yeniden oluşturuluyor...")
                success = self._replace_via_text_recreation(input_path, entities, output_path)
                if success:
                    return True
            
            self.logger.error("Tüm değiştirme yöntemleri başarısız")
            return False
            
        except Exception as e:
            self.logger.error(f"PDF değiştirme kritik hatası: {e}")
            return False
    
    def _replace_via_docx_method(self, input_path: str, entities: List[Dict], 
                               output_path: str, progress) -> bool:
        """DOCX dönüştürme yöntemi"""
        try:
            from pdf2docx import parse
            import tempfile
            
            with tempfile.TemporaryDirectory() as temp_dir:
                # 1) PDF -> DOCX
                docx_path = os.path.join(temp_dir, "temp.docx")
                progress(0.55, desc="PDF DOCX'e dönüştürülüyor...")
                parse(input_path, docx_path, pages=None)
                
                if not os.path.exists(docx_path):
                    raise Exception("PDF->DOCX dönüştürme başarısız")
                
                # 2) DOCX içinde metinleri değiştir
                progress(0.6, desc="DOCX'te metinler değiştiriliyor...")
                replacements_made = self._replace_text_in_docx_advanced(docx_path, entities)
                if replacements_made == 0:
                    self.logger.warning("DOCX'te hiçbir değiştirme yapılamadı")
                    return False
                
                # 3) DOCX -> PDF (çoklu yöntem)
                progress(0.7, desc="DOCX PDF'e dönüştürülüyor...")
                success = self._convert_docx_to_pdf_robust(docx_path, output_path)
                if success:
                    self.logger.info(f"DOCX yöntemi başarılı: {replacements_made} değiştirme")
                    return True
                return False
        except Exception as e:
            self.logger.error(f"DOCX yöntemi hatası: {e}")
            return False
    
    def _replace_text_in_docx_advanced(self, docx_path: str, entities: List[Dict]) -> int:
        """DOCX'te gelişmiş metin değiştirme"""
        try:
            from docx import Document
            import re
            
            doc = Document(docx_path)
            total_replacements = 0
            
            # Replacement map
            replacement_map = {}
            for entity in entities:
                original = entity.get('word', '').strip()
                replacement = entity.get('replacement', original).strip()
                if original and replacement != original:
                    replacement_map[original] = replacement
            
            if not replacement_map:
                return 0
            
            self.logger.info(f"DOCX'te değiştirilecek (örnek): {list(replacement_map.keys())[:5]}...")
            
            # Paragraflar
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    replaced = self._replace_in_paragraph_advanced(paragraph, replacement_map)
                    total_replacements += replaced
            
            # Tablolar
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                replaced = self._replace_in_paragraph_advanced(paragraph, replacement_map)
                                total_replacements += replaced
            
            # Header/Footer
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        total_replacements += self._replace_in_paragraph_advanced(paragraph, replacement_map)
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        total_replacements += self._replace_in_paragraph_advanced(paragraph, replacement_map)
            
            doc.save(docx_path)
            return total_replacements
        except Exception as e:
            self.logger.error(f"DOCX metin değiştirme hatası: {e}")
            return 0
    
    def _replace_in_paragraph_advanced(self, paragraph, replacement_map: Dict[str, str]) -> int:
        """Paragrafta gelişmiş metin değiştirme"""
        import re
        if not paragraph.text.strip():
            return 0
        original_text = paragraph.text
        new_text = original_text
        replacements_count = 0
        
        for original, replacement in replacement_map.items():
            if original in new_text:
                pattern = r'\b' + re.escape(original) + r'\b'
                matches = len(re.findall(pattern, new_text))
                if matches > 0:
                    new_text = re.sub(pattern, replacement, new_text)
                    replacements_count += matches
                    self.logger.debug(f"'{original}' -> '{replacement}' ({matches} kez)")
        
        if new_text != original_text:
            # Run'ları temizleyip yeni metni ekle
            for r in paragraph.runs:
                r.text = ""
            paragraph.text = new_text
        return replacements_count
    
    def _convert_docx_to_pdf_robust(self, docx_path: str, output_path: str) -> bool:
        """Güçlü DOCX->PDF dönüştürme (3 yöntem)"""
        conversions = [
            ('docx2pdf', self._convert_with_docx2pdf),
            ('libreoffice', self._convert_with_libreoffice),
            ('comtypes', self._convert_with_comtypes)
        ]
        for method_name, func in conversions:
            try:
                self.logger.info(f"{method_name} ile dönüştürme deneniyor...")
                if func(docx_path, output_path):
                    self.logger.info(f"{method_name} başarılı!")
                    return True
            except Exception as e:
                self.logger.warning(f"{method_name} başarısız: {e}")
        return False
    
    def _convert_with_docx2pdf(self, docx_path: str, output_path: str) -> bool:
        try:
            from docx2pdf import convert
            convert(docx_path, output_path)
            return os.path.exists(output_path) and os.path.getsize(output_path) > 1000
        except Exception:
            return False
    
    def _convert_with_libreoffice(self, docx_path: str, output_path: str) -> bool:
        try:
            import subprocess
            outdir = os.path.dirname(output_path) or "."
            cmd = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', outdir, docx_path]
            result = subprocess.run(cmd, capture_output=True, timeout=120)
            if result.returncode == 0:
                expected = os.path.join(outdir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
                if os.path.exists(expected):
                    if expected != output_path:
                        shutil.move(expected, output_path)
                    return True
            return False
        except Exception:
            return False
    
    def _convert_with_comtypes(self, docx_path: str, output_path: str) -> bool:
        try:
            import comtypes.client
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(docx_path))
            doc.SaveAs(os.path.abspath(output_path), FileFormat=17)
            doc.Close()
            word.Quit()
            return os.path.exists(output_path) and os.path.getsize(output_path) > 1000
        except Exception:
            return False
    
    def _replace_via_pymupdf_method(self, input_path: str, entities: List[Dict], output_path: str) -> bool:
        """PyMuPDF ile direkt PDF değiştirme"""
        try:
            import fitz
            doc = fitz.open(input_path)
            total_replacements = 0
            
            replacement_map = {}
            for entity in entities:
                original = entity.get('word', '').strip()
                replacement = entity.get('replacement', original).strip()
                if original and replacement != original:
                    replacement_map[original] = replacement
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                for original, replacement in replacement_map.items():
                    text_instances = page.search_for(original)
                    for rect in text_instances:
                        page.add_redact_annot(rect, fill=(1, 1, 1))
                        page.insert_text(
                            rect[:2],  # sol üst
                            replacement,
                            fontsize=10,
                            color=(0, 0, 0)
                        )
                        total_replacements += 1
                page.apply_redactions()
            
            doc.save(output_path)
            doc.close()
            self.logger.info(f"PyMuPDF: {total_replacements} değiştirme yapıldı")
            return total_replacements > 0
        except Exception as e:
            self.logger.error(f"PyMuPDF hatası: {e}")
            return False
    
    def _replace_via_text_recreation(self, input_path: str, entities: List[Dict], output_path: str) -> bool:
        """Metin çıkarıp yeniden oluşturma yöntemi (basit)"""
        try:
            import fitz
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
            
            doc = fitz.open(input_path)
            c = canvas.Canvas(output_path, pagesize=A4)
            
            replacement_map = {}
            for entity in entities:
                original = entity.get('word', '').strip()
                replacement = entity.get('replacement', original).strip()
                if original and replacement != original:
                    replacement_map[original] = replacement
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                for original, replacement in replacement_map.items():
                    text = text.replace(original, replacement)
                
                c.showPage()
                textobject = c.beginText(50, 780)
                for line in text.split('\n'):
                    textobject.textLine(line[:120])
                c.drawText(textobject)
            
            c.save()
            doc.close()
            return os.path.exists(output_path) and os.path.getsize(output_path) > 1000
        except Exception as e:
            self.logger.error(f"Text recreation hatası: {e}")
            return False
    
    def validate_pdf_changes(self, output_path: str, entities: List[Dict]) -> Dict:
        """PDF'teki değişiklikleri doğrula"""
        try:
            import fitz
            doc = fitz.open(output_path)
            full_text = ""
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                full_text += page.get_text()
            doc.close()
            
            validation_result = {
                'total_entities': len(entities),
                'successfully_replaced': 0,
                'still_present': [],
                'replacement_found': [],
                'partial_replacements': []
            }
            for entity in entities:
                original = entity.get('word', '').strip()
                replacement = entity.get('replacement', original).strip()
                if original in full_text:
                    validation_result['still_present'].append(original)
                else:
                    validation_result['successfully_replaced'] += 1
                if replacement in full_text and replacement != original:
                    validation_result['replacement_found'].append(replacement)
            return validation_result
        except Exception as e:
            self.logger.error(f"Validasyon hatası: {e}")
            return {'error': str(e)}
    
    def format_debug_info(self, validation_result: Dict, entities: List[Dict]) -> str:
        """Debug bilgilerini formatla"""
        if 'error' in validation_result:
            return f"❌ Validasyon hatası: {validation_result['error']}"
        
        debug_info = f"""
## 🔍 Debug Bilgileri

### Değiştirme Sonuçları:
- **Toplam Entity**: {validation_result['total_entities']}
- **Başarılı Değiştirme**: {validation_result['successfully_replaced']}
- **Hala Mevcut**: {len(validation_result.get('still_present', []))}
- **Yeni Değerler Bulundu**: {len(validation_result.get('replacement_found', []))}

### Hala Mevcut Metinler:
{validation_result.get('still_present', [])[:5]}

### Bulunan Yeni Değerler:
{validation_result.get('replacement_found', [])[:5]}
        """.strip()
        return debug_info
    
    def create_detailed_statistics(self, entities: List[Dict], validation_result: Dict) -> pd.DataFrame:
        """Detaylı istatistik tablosu"""
        if not entities:
            return pd.DataFrame(columns=['Veri Türü', 'Tespit', 'Değiştirildi', 'Başarı Oranı', 'Örnek'])
        
        stats = {}
        for entity in entities:
            etype = entity['entity']
            if etype not in stats:
                stats[etype] = {'detected': 0, 'examples': [], 'originals': [], 'replacements': []}
            stats[etype]['detected'] += 1
            stats[etype]['originals'].append(entity.get('word', ''))
            stats[etype]['replacements'].append(entity.get('replacement', ''))
            if len(stats[etype]['examples']) < 2:
                stats[etype]['examples'].append(entity.get('word', ''))
        
        rows = []
        names = {
            'ad_soyad': '👤 Ad Soyad',
            'tc_kimlik': '🆔 TC Kimlik', 
            'telefon': '📱 Telefon',
            'adres': '📍 Adres',
            'para': '💰 Para',
            'tarih': '📅 Tarih',
            'email': '📧 E-posta',
            'sirket': '🏢 Şirket'
        }
        for etype, data in stats.items():
            successful = 0
            for original in data['originals']:
                if original not in validation_result.get('still_present', []):
                    successful += 1
            success_rate = (successful / data['detected'] * 100) if data['detected'] > 0 else 0
            rows.append({
                'Veri Türü': names.get(etype, etype.title()),
                'Tespit': data['detected'],
                'Değiştirildi': successful,
                'Başarı Oranı': f"{success_rate:.1f}%",
                'Örnek': ', '.join(data['examples'][:2])
            })
        return pd.DataFrame(rows)
    
    def generate_detailed_summary(self, entities: List[Dict], validation_result: Dict, filename: str) -> str:
        """Detaylı özet oluştur"""
        if not entities:
            return "❌ Hiçbir kişisel bilgi tespit edilmedi."
        
        total_entities = len(entities)
        successful = validation_result.get('successfully_replaced', 0)
        still_present = len(validation_result.get('still_present', []))
        success_rate = (successful / total_entities * 100) if total_entities > 0 else 0
        
        if success_rate >= 90:
            status_emoji = "🟢"; status_text = "Mükemmel"
        elif success_rate >= 70:
            status_emoji = "🟡"; status_text = "İyi"
        else:
            status_emoji = "🔴"; status_text = "Dikkat Gerekli"
        
        summary = f"""
## 📊 Detaylı İşlem Raporu

**Dosya:** {filename}  
**Tarih:** {datetime.now().strftime('%d.%m.%Y %H:%M')}

### {status_emoji} Genel Başarı: {status_text} ({success_rate:.1f}%)

### 🎯 Değiştirme Sonuçları
- **Toplam Tespit:** {total_entities} adet
- **Başarılı Değiştirme:** {successful} adet  
- **Hala Mevcut:** {still_present} adet

### 📈 Performans Analizi
""".strip()
        if success_rate >= 90:
            summary += "\n\n✅ **Harika!** Neredeyse tüm kişisel bilgiler başarıyla değiştirildi."
        elif success_rate >= 70:
            summary += "\n\n⚠️ **İyi performans** ancak bazı bilgiler değiştirilemedi. Manuel kontrol önerilir."
        else:
            summary += "\n\n🚨 **Dikkat!** Çoğu bilgi değiştirilemedi. Farklı yöntem denemeyi düşünün."
        
        if still_present > 0:
            still_present_list = validation_result.get('still_present', [])[:3]
            summary += "\n\n### 🔍 Değiştirilemeyen Örnekler:\n"
            for item in still_present_list:
                summary += f"- `{item}`\n"
        return summary
    
    def log_processing_with_validation(self, filename: str, entities: List[Dict], validation_result: Dict):
        """Validasyon sonuçlarıyla loglama"""
        log_entry = {
            'timestamp': datetime.now().isoformat(),
            'filename': filename,
            'total_entities': len(entities),
            'successful_replacements': validation_result.get('successfully_replaced', 0),
            'still_present_count': len(validation_result.get('still_present', [])),
            'success_rate': (validation_result.get('successfully_replaced', 0) / len(entities) * 100) if entities else 0,
            'entities': entities,
            'validation_result': validation_result
        }
        log_file = f"logs/detailed_session_{datetime.now().strftime('%Y%m%d')}.json"
        logs = []
        if os.path.exists(log_file):
            try:
                with open(log_file, 'r', encoding='utf-8') as f:
                    logs = json.load(f)
            except Exception:
                logs = []
        logs.append(log_entry)
        with open(log_file, 'w', encoding='utf-8') as f:
            json.dump(logs, f, ensure_ascii=False, indent=2)
    
    def get_processing_history(self) -> pd.DataFrame:
        """İşlem geçmişini al"""
        try:
            log_files = []
            for i in range(7):
                date_str = (datetime.now() - pd.Timedelta(days=i)).strftime('%Y%m%d')
                log_file = f"logs/detailed_session_{date_str}.json"
                if os.path.exists(log_file):
                    log_files.append(log_file)
            all_logs = []
            for log_file in log_files:
                with open(log_file, 'r', encoding='utf-8') as f:
                    logs = json.load(f)
                    all_logs.extend(logs)
            if not all_logs:
                return pd.DataFrame(columns=['Tarih', 'Dosya', 'Toplam Tespit', 'Başarılı', 'Başarı Oranı'])
            rows = []
            for log in all_logs[-50:]:
                timestamp = datetime.fromisoformat(log['timestamp'])
                rows.append({
                    'Tarih': timestamp.strftime('%d.%m.%Y %H:%M'),
                    'Dosya': log['filename'],
                    'Toplam Tespit': log['total_entities'],
                    'Başarılı': log.get('successful_replacements', 0),
                    'Başarı Oranı': f"{log.get('success_rate', 0):.1f}%"
                })
            return pd.DataFrame(rows)
        except Exception:
            return pd.DataFrame(columns=['Tarih', 'Dosya', 'Toplam Tespit', 'Başarılı', 'Başarı Oranı'])
    
    # -----------------------------
    # Gradio UI
    # -----------------------------
    def create_enhanced_interface(self):
        """Gelişmiş Gradio arayüzü"""
        css = """
        .gradio-container { font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; }
        .debug-info { background-color: #f8f9fa; padding: 10px; border-radius: 5px; font-family: monospace; font-size: 12px; }
        """
        
        with gr.Blocks(css=css, title="PDF Kişisel Bilgi Anonimleştirici - Gelişmiş", theme=gr.themes.Soft()) as interface:
            gr.Markdown("""
            # 🔐 PDF Kişisel Bilgi Anonimleştirici — Gelişmiş
            Yüklediğiniz PDF'teki kişisel bilgileri **tespit eder** ve **gerçekten değiştirir**.  
            Sonuçları doğrular, rapor ve istatistik üretir.
            """)
            
            with gr.Tabs():
                # ------------------ Ana İşlem Sekmesi ------------------
                with gr.TabItem("📄 PDF İşleme", elem_id="main-tab"):
                    with gr.Row():
                        with gr.Column(scale=1):
                            pdf_input = gr.File(
                                label="📎 PDF Dosyası Yükleyin",
                                file_types=[".pdf"],
                                type="filepath"
                            )
                            
                            gr.Markdown("### ⚙️ İşlem Ayarları")
                            
                            confidence_threshold = gr.Slider(
                                minimum=0.1, maximum=1.0, value=0.7, step=0.1,
                                label="🎯 Güven Eşiği"
                            )
                            
                            replacement_strategy = gr.Radio(
                                choices=["Maskeleme (*)", "Sahte Değerler", "Genel Değerler"],
                                value="Sahte Değerler",
                                label="🔁 Değiştirme Stratejisi"
                            )
                            
                            with gr.Accordion("🧩 Özel Değer Haritası (JSON)", open=False):
                                gr.Markdown("Örnek:")
                                example_json = {
                                    "ad_soyad": "Mehmet Yılmaz",
                                    "telefon": "05551234567",
                                    "email": "anonim@example.com",
                                    "adres": "Ankara, Türkiye"
                                }
                                gr.Code(value=json.dumps(example_json, ensure_ascii=False, indent=2), language="json")
                                custom_replacements = gr.Textbox(
                                    label="Özel Değerler (JSON)",
                                    placeholder='{"ad_soyad": "Mehmet Yılmaz", "telefon": "0555..."}',
                                    lines=5
                                )
                            
                            conversion_method = gr.Dropdown(
                                choices=["auto", "DOCX Yöntemi", "PyMuPDF Direkt", "Text Recreation"],
                                value="auto",
                                label="🛠️ Dönüştürme / Değiştirme Yöntemi"
                            )
                            
                            enable_logging = gr.Checkbox(value=True, label="🧾 Log kaydı tut")
                            enable_debug = gr.Checkbox(value=False, label="🐞 Debug bilgilerini göster")
                            
                            process_btn = gr.Button("🚀 İşlemi Başlat", variant="primary")
                        
                        with gr.Column(scale=1):
                            output_pdf = gr.File(label="📤 Anonimleştirilmiş PDF", file_count="single")
                            status_text = gr.Markdown("⏳ Henüz işlem yapılmadı.")
                            summary_md = gr.Markdown("")
                            stats_df = gr.Dataframe(headers=["Veri Türü", "Tespit", "Değiştirildi", "Başarı Oranı", "Örnek"], interactive=False)
                            debug_out = gr.Textbox(label="Debug", lines=10, visible=False)
                    
                    # Buton-click bağlama
                    def _wrap_process(pdf, thr, strat, custom_json, log_on, dbg_on, method, progress=gr.Progress()):
                        out_path, status, summary, df, dbg = self.process_pdf_with_real_replacement(
                            pdf, thr, strat, custom_json or "", log_on, dbg_on, method, progress
                        )
                        # Debug görünürlüğü
                        return (
                            out_path, 
                            status, 
                            summary, 
                            df, 
                            gr.update(value=dbg, visible=bool(dbg_on))
                        )
                    
                    process_btn.click(
                        _wrap_process,
                        inputs=[pdf_input, confidence_threshold, replacement_strategy, custom_replacements, enable_logging, enable_debug, conversion_method],
                        outputs=[output_pdf, status_text, summary_md, stats_df, debug_out],
                        api_name="process_pdf"
                    )
                
                # ------------------ Geçmiş ------------------
                with gr.TabItem("🗂️ Geçmiş"):
                    gr.Markdown("Son işlemlerin özeti (loglardan derlenir).")
                    history_df = gr.Dataframe(interactive=False)
                    refresh_btn = gr.Button("🔄 Yenile")
                    
                    def _load_history():
                        return self.get_processing_history()
                    
                    refresh_btn.click(_load_history, inputs=None, outputs=history_df)
                    # Sayfa açılır açılmaz doldurmak için:
                    interface.load(_load_history, inputs=None, outputs=history_df)
                
                # ------------------ Yardım ------------------
                # ------------------ Yardım ------------------
                with gr.TabItem("❓ Yardım"):
                    gr.Markdown("""
### Sık İpuçları
- **PyMuPDF Direkt** hızlıdır ama karmaşık PDF'lerde her metni bulamayabilir.
- **DOCX Yöntemi** daha isabetli olabilir; sisteminizde `libreoffice` ya da `Word`/`comtypes` yoksa `docx2pdf` ile de çalışır.
- **Text Recreation** sayfa düzenini korumaz; en basit/garanti seçenektir, rapor için uygundur.

### Gerekli Paketler
```bash
pip install gradio pandas pymupdf pdf2docx python-docx docx2pdf reportlab
# (İsteğe bağlı) libreoffice veya comtypes (Windows)
markdown
Copy
Edit
                """)

    # Tabs kapandı, Blocks kapanmadan önce interface'i döndürüyoruz
                return interface
if __name__ == "__main__":
    app = EnhancedAnonymizationApp()
    demo = app.create_enhanced_interface()
# local/lan için ayarları ihtiyaçlarına göre düzenleyebilirsin
    demo.launch(server_name="0.0.0.0", server_port=7860, show_error=True)
 
